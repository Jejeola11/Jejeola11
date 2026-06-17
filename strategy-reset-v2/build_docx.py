#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x0B, 0x25, 0x45)
GOLD = RGBColor(0xC9, 0xA2, 0x4B)
DARK = RGBColor(0x11, 0x22, 0x3A)
GREY = RGBColor(0x6B, 0x79, 0x90)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
NAVY_HEX = "0B2545"
GOLD_HEX = "C9A24B"
LIGHT_HEX = "F6F8FB"

def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear'); sh.set(qn('w:color'), 'auto'); sh.set(qn('w:fill'), hex_color)
    tcPr.append(sh)

def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement('w:tcMar')
    for tag, val in (('top',top),('bottom',bottom),('start',left),('end',right)):
        e = OxmlElement(f'w:{tag}'); e.set(qn('w:w'), str(val)); e.set(qn('w:type'),'dxa'); m.append(e)
    tcPr.append(m)

def no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        e = OxmlElement(f'w:{edge}'); e.set(qn('w:val'),'none'); e.set(qn('w:sz'),'0'); borders.append(e)
    tblPr.append(borders)

def run(p, text, size=11, bold=False, color=DARK, italic=False):
    r = p.add_run(text); r.font.size = Pt(size); r.font.bold = bold; r.font.italic = italic
    r.font.color.rgb = color; r.font.name = 'Calibri'
    return r

def heading(doc, text, size=17):
    p = doc.add_paragraph(); p.space_before = Pt(14)
    r = run(p, text.upper(), size=size, bold=True, color=NAVY)
    pf = p.paragraph_format; pf.space_before = Pt(16); pf.space_after = Pt(4)
    # gold bottom border
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'12'); bottom.set(qn('w:space'),'3'); bottom.set(qn('w:color'),GOLD_HEX)
    pbdr.append(bottom); pPr.append(pbdr)
    return p

def sub(doc, text):
    p = doc.add_paragraph(); pf = p.paragraph_format; pf.space_before = Pt(8); pf.space_after = Pt(2)
    run(p, text, size=12, bold=True, color=RGBColor(0x1C,0x35,0x59))
    return p

def body(doc, text, size=11, color=DARK, after=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    run(p, text, size=size, color=color)
    return p

def bullet(doc, runs):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(2)
    for t,b in runs:
        run(p, t, size=11, bold=b, color=DARK)
    return p

def small(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(6)
    run(p, text, size=8.5, color=GREY, italic=True)
    return p

def make_table(doc, rows, header=True, widths=None, total_row=False, num_cols=None):
    t = doc.add_table(rows=0, cols=len(rows[0]))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, rowdata in enumerate(rows):
        cells = t.add_row().cells
        is_header = header and i == 0
        is_total = total_row and i == len(rows)-1
        for j, val in enumerate(rowdata):
            c = cells[j]; set_cell_margins(c)
            c.text = ''
            para = c.paragraphs[0]; para.paragraph_format.space_after = Pt(0)
            if num_cols and j in num_cols:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            if is_header:
                shade(c, NAVY_HEX); run(para, str(val), size=9.5, bold=True, color=WHITE)
            elif is_total:
                shade(c, GOLD_HEX); run(para, str(val), size=10, bold=True, color=NAVY)
            else:
                if i % 2 == 0: shade(c, LIGHT_HEX)
                run(para, str(val), size=9.5, color=DARK)
    if widths:
        for row in t.rows:
            for idx, w in enumerate(widths):
                row.cells[idx].width = Inches(w)
    no_borders(t)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# ---------------- BUILD CLIENT PROPOSAL ----------------
doc = Document()
sec = doc.sections[0]
sec.top_margin = Cm(1.6); sec.bottom_margin = Cm(1.6); sec.left_margin = Cm(1.8); sec.right_margin = Cm(1.8)
style = doc.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)

# COVER (navy block via single-cell table)
cover = doc.add_table(rows=1, cols=1); cover.alignment = WD_TABLE_ALIGNMENT.CENTER
cc = cover.rows[0].cells[0]; shade(cc, NAVY_HEX); set_cell_margins(cc, top=240, bottom=240, left=260, right=260)
cc.text = ''
p = cc.paragraphs[0]; p.paragraph_format.space_after = Pt(4)
run(p, 'OFFICIALLY · INVESTED — ANYA MEHRA', size=10, bold=True, color=GOLD)
p2 = cc.add_paragraph(); p2.paragraph_format.space_after = Pt(2)
run(p2, 'Strategy Reset v2', size=30, bold=True, color=WHITE)
p3 = cc.add_paragraph(); p3.paragraph_format.space_after = Pt(8)
run(p3, 'New direction, the investment to run it, and the milestone close-out', size=13, bold=True, color=GOLD)
p4 = cc.add_paragraph()
run(p4, 'Prepared for the client call  ·  17 June 2026  ·  Covers: Milestone 1 release · the new content engine · production & ad budget · restructured performance bonus · payment options', size=9, color=RGBColor(0xCD,0xD8,0xE6))
no_borders(cover)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

body(doc, "The first 15 days are delivered and the system is built. They didn't break through on reach — so this document does two things: it closes out the completed milestone, and lays out a sharper, evidence-led engine for the next cycle, with a clear, fixed investment so we both know exactly what we're putting in before we set off.", size=11.5, color=RGBColor(0x33,0x45,0x5F), after=4)

# 1 Milestone
heading(doc, '1 · Milestone 1 — First 15 Days (Delivered)')
body(doc, "The first 15-day content block — scripts, AI avatar production, platform-ready edits and scheduling across Instagram, TikTok and YouTube — is complete and delivered. This is contracted, finished work and is independent of the new direction below.")
make_table(doc, [
    ['Deliverable','Amount'],
    ['First 15 days — full content block, produced & published (completed)','$750.00'],
    ['Release to close Milestone 1','$750.00'],
], widths=[5.6,1.4], total_row=True, num_cols=[1])
small(doc, "Recommended action: release Milestone 1 today. The new cycle starts on a clean slate once this is settled.")

# 2 Why
heading(doc, "2 · Why Month One Didn't Break Through")
body(doc, "An honest read — so the fix is targeted, not guesswork:")
bullet(doc, [("The hook lost the first 3 seconds. ", True), ("Static opens get scrolled past before the message lands.", False)])
bullet(doc, [("We led with polish, not friction. ", True), ("The feed rewards a sharp, contrarian opening line, not a calm intro.", False)])
bullet(doc, [("No proven structure underneath the posts. ", True), ("Good content, but not shaped like what already wins in this niche.", False)])
bullet(doc, [("No paid reach. ", True), ("Organic-only on cold accounts barely distributes. Nothing was filling the funnel.", False)])

# 3 New direction
heading(doc, '3 · The New Direction — Model What Already Wins, Then Put Spend Behind It')
sub(doc, 'Step 1 — Reverse-engineer proven winners')
body(doc, "We study posts in the business-acquisition / “boring wealth” niche that have already done millions of views, and break down the structure that made them work — hook, pacing, on-screen text, the emotional turn, the call to action. We model that proven structure, then write original Anya content into it. The framework is borrowed; the script, voice and brand stay 100% ours.")
sub(doc, 'Step 2 — Build for all three platforms')
body(doc, "Each idea is produced as native content for YouTube, Instagram and TikTok — correct format, length and end-card per platform — using the AI avatar + motion stack. One proven structure becomes three platform-native assets: more surface area for the same production effort.")
sub(doc, 'Step 3 — Paid boost: the reach lever (7-day push)')
body(doc, "Once a post is live and performing, we put paid spend behind the best ones for a full week across all three platforms. This is what gets Anya in front of large, cold audiences fast — the difference between a good post nobody sees and a good post that compounds followers and signups. Organic earns the right; paid pours fuel on it.")

# 4 Budget
heading(doc, '4 · Production & Ad Budget — Everything It Takes To Run One Cycle')
body(doc, "Full transparency on what the engine costs before we start. Two parts: the monthly tool stack that produces the content, and the one-week paid push that distributes it.")
sub(doc, 'A · Monthly tool stack (production)')
make_table(doc, [
    ['Tool','What it does','Plan','Monthly'],
    ['HeyGen','AI avatar video + Seedance motion (the talking, moving Anya)','Pro / Scale','$99'],
    ['Higgsfield','Cinematic B-roll & motion shots for the first 3 seconds','Ultimate','$49'],
    ['Claude','Scripts, contrarian hooks, captions, structure breakdowns','Pro','$20'],
    ['Metricool','Scheduling + analytics across all 3 platforms','Advanced','$22'],
    ['ElevenLabs','Voiceover / narration','Creator','$22'],
    ['CapCut Pro','Editing, captions, end-cards','Pro','$10'],
    ['Canva Pro','Carousels, covers, thumbnails','Pro','$13'],
    ['Tool stack subtotal — per month','','','$235'],
], widths=[1.1,3.3,1.0,0.9], total_row=True, num_cols=[3])
sub(doc, 'B · Paid boost (7-day push across 3 platforms)')
make_table(doc, [
    ['Platform','Format boosted','Per day','Days','Total'],
    ['Instagram','Reels + carousels','$30','7','$210'],
    ['TikTok','Spark Ads / Promote','$30','7','$210'],
    ['YouTube','Shorts / in-feed','$25','7','$175'],
    ['Paid boost subtotal — recommended (Standard) tier','','','','$595'],
], widths=[1.1,2.6,0.8,0.6,0.9], total_row=True, num_cols=[2,3,4])
small(doc, "Boost is adjustable:  Lean ≈ $315/week (~$15/day each)  ·  Standard (recommended) ≈ $595/week  ·  Aggressive ≈ $1,050/week (~$50/day each). Standard is the level that reliably moves follower and signup numbers in a 7-day window.")
sub(doc, 'C · One-cycle total')
make_table(doc, [
    ['Component','Amount'],
    ['Monthly tool stack (production)','$235'],
    ['Paid boost — 7-day push (Standard)','$595'],
    ['Contingency buffer (~10% — failed renders, platform ad fees, re-exports)','$85'],
    ['Production & ad budget — one cycle','$915'],
], widths=[5.6,1.4], total_row=True, num_cols=[1])
small(doc, "This is the project's running cost, separate from the Milestone 1 release. Any tools already owned can be deducted directly.")

# 5 Bonus
heading(doc, '5 · Restructured Performance Bonus — Tied To What We Control')
body(doc, "The original bonus — $500 for the first 40 paid signups — rewards an outcome that sits past our handoff. Whether a warm lead converts to a paid challenge depends on the sales page, the price and the follow-up — none of which we run. So we're proposing to re-tie the same reward to the metric we do own and drive directly:")
sub(doc, 'New bonus: paid on free-community signups')
make_table(doc, [
    ['Structure','Rate','Target','Bonus'],
    ['Per verified member into the free community','$5 / signup','100 members','$500'],
    ['Same budget as before — now tied to a number we can move','','','$500'],
], widths=[3.4,1.3,1.3,0.9], total_row=True, num_cols=[2,3])
body(doc, "Why this is better for you: getting people into the free community is exactly the job our content does — so the incentive lines up with the work, and it builds you an owned audience you keep marketing to for free. It's fully measurable, fully attributable, and scalable on your call: if you want more momentum, we lift the target (e.g. 250 members → $1,000) at the same per-signup rate.")

# 6 Test
heading(doc, '6 · The 10-Day Test & Decision')
body(doc, "We run the new engine on 10 scripts over 10 days, with the 7-day paid push behind the best performers. At the end we read the numbers together — reach, follower growth, and free-community signups — and decide whether to scale, adjust, or stop. No open-ended commitment; a clear checkpoint.")

# 7 Ask + payment
heading(doc, '7 · What We’re Asking For Today & How To Fund It')
make_table(doc, [
    ['#','Item','Amount'],
    ['1','Milestone 1 — first 15 days (delivered)','$750'],
    ['2','Production & ad budget — one cycle (tools + 7-day boost)','$915'],
    ['3','Performance bonus — paid on free-community signups (on results)','$500'],
], widths=[0.4,5.2,1.4], num_cols=[2])
small(doc, "Item 3 is paid only on delivered signups. Items 1 & 2 are what we need to close out and set off.")
sub(doc, 'Two ways to pay')
make_table(doc, [
    ['Option A — Release via milestone','Option B — Direct transfer (fastest)'],
    ['✓  Funds settle in ~7 days\n✓  Stays inside the platform’s protection for both sides\n\nBest if you’d rather keep everything on the existing milestone rails.',
     '✓  Sent to my personal US account\n✓  Clears in ~2 days — we start almost immediately\n\nBest if you want the new cycle moving this week. Off-platform, so it’s a direct trust arrangement between us.'],
], header=False, widths=[3.5,3.5])
small(doc, "Either route works on my end — it's purely a question of how fast you want the next cycle to start.")

# Footer
fp = doc.add_paragraph(); fp.paragraph_format.space_before = Pt(14)
pPr = fp._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr'); top = OxmlElement('w:top')
top.set(qn('w:val'),'single'); top.set(qn('w:sz'),'6'); top.set(qn('w:space'),'4'); top.set(qn('w:color'),'E2E8F1')
pbdr.append(top); pPr.append(pbdr)
run(fp, "Officially · Invested — Anya Mehra  |  Strategy Reset v2  |  Prepared 17 June 2026. All figures are estimates for one test cycle and adjustable on the call. No income or results are guaranteed; projections are honest planning numbers, not promises.", size=8, color=GREY, italic=True)

doc.save('/home/user/Jejeola11/strategy-reset-v2/Officially_Invested_Strategy_Reset_v2_Client_Proposal.docx')
print('CLIENT PROPOSAL saved')

# ---------------- BUILD RUN-SHEET ----------------
d2 = Document()
s2 = d2.sections[0]; s2.top_margin = Cm(1.6); s2.bottom_margin=Cm(1.6); s2.left_margin=Cm(1.8); s2.right_margin=Cm(1.8)
d2.styles['Normal'].font.name='Calibri'; d2.styles['Normal'].font.size=Pt(11)
tp = d2.add_paragraph(); run(tp,'Call Run-Sheet — Strategy Reset v2', size=22, bold=True, color=NAVY)
sp = d2.add_paragraph(); run(sp,'INTERNAL · for you only', size=11, bold=True, color=GOLD)
gp = d2.add_paragraph(); gp.paragraph_format.space_after=Pt(10)
run(gp,'Goal of the call: (1) release the $750 milestone, (2) get sign-off on the new direction, (3) lock the production budget, (4) swap the bonus to free-community signups, (5) agree how they pay. Keep the tone calm and confident — you delivered the work. This is a better plan, not an apology.', size=10.5, color=DARK)

def quote(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.25); p.paragraph_format.space_after=Pt(6)
    pPr=p._p.get_or_add_pPr(); pbdr=OxmlElement('w:pBdr'); left=OxmlElement('w:left')
    left.set(qn('w:val'),'single'); left.set(qn('w:sz'),'18'); left.set(qn('w:space'),'8'); left.set(qn('w:color'),GOLD_HEX)
    pbdr.append(left); pPr.append(pbdr)
    run(p, text, size=10.5, italic=True, color=RGBColor(0x1C,0x35,0x59))

heading(d2,'0 · Open (30 seconds)', size=14)
quote(d2,'“Thanks for making time. The first 15 days are built and delivered. They didn’t break through on reach — and I’ve got a clear, honest read on why, plus a sharper plan that fixes it. I’ll walk you through the new direction, exactly what it costs to run, and a cleaner way to structure the bonus.”')
body(d2,'Share the proposal document.', size=10.5)

heading(d2,'1 · Close out Milestone 1 — $750 (do this first)', size=14)
bullet(d2,[('The first 15 days are complete, contracted work — independent of whatever we decide next.',False)])
bullet(d2,[('Ask plainly: ',False),('“Can we release the $750 for the first block today so we start the next cycle clean?”',True)])
bullet(d2,[('If they hesitate on results: that milestone is for production & delivery, which is done. The reach problem is exactly what the new plan solves — and they’re not funding that blindly; the whole budget is laid out.',False)])

heading(d2,'2 · Why it underperformed (60 seconds, matter-of-fact)', size=14)
for t in ['Hook lost the first 3 seconds (static opens).','Led with polish, not a contrarian opening line.','No proven structure under the posts.','No paid reach — organic-only on cold accounts barely distributes.']:
    bullet(d2,[(t,False)])
quote(d2,'“None of these are fatal. They’re all fixable, and the fix is the new direction.”')

heading(d2,'3 · The new direction (core pitch)', size=14)
bullet(d2,[('Model what already wins: ',True),('find niche posts doing millions of views, break down the structure, write original Anya content into that proven shape.',False)])
bullet(d2,[('If he worries about copying: ',True),('“We borrow the structure, not the content. Script, voice and brand stay 100% ours. Every top creator works this way — they reuse what the algorithm already rewards.”',False)])
bullet(d2,[('Three platforms native: ',True),('one structure → YouTube + Instagram + TikTok.',False)])
bullet(d2,[('Paid boost, 7-day push: ',True),('“This is the lever that was missing. The difference between a good post nobody sees and one that compounds followers and signups.”',False)])

heading(d2,'4 · The budget (lead with the total)', size=14)
bullet(d2,[('Tool stack: $235/mo ',True),('— HeyGen, Higgsfield, Claude, Metricool, ElevenLabs, CapCut, Canva.',False)])
bullet(d2,[('Paid boost: $595 ',True),('for a 7-day push (recommend Standard; Lean ~$315 / Aggressive ~$1,050).',False)])
bullet(d2,[('+10% buffer → one-cycle total ≈ $915. ',True),('Anything he already pays for, we deduct.',False)])

heading(d2,'5 · Restructure the bonus (your key ask)', size=14)
quote(d2,'“The old bonus was $500 for the first 40 paid signups. Problem: whether a lead converts to paid depends on the sales page, price and follow-up — things I don’t control. So I’d be paid on something outside my hands, which isn’t fair to either of us.”')
quote(d2,'“What I do control is getting people into the free community. So I want to tie the same $500 to that — $5 per verified free-community member, target 100. Same budget, but now it’s tied to a number I can actually move — and it builds you an owned audience you keep marketing to for free.”')
bullet(d2,[('Scalable hook: ',True),('“Want more momentum? We just raise the target at the same rate — 250 members for $1,000. You decide the dial.”',False)])
bullet(d2,[('If he pushes on rate: ',True),('$5/signup is the anchor. Flex the target up, not the rate down.',False)])

heading(d2,'6 · The test & decision (lowers his risk)', size=14)
quote(d2,'“We run 10 scripts over 10 days with the paid push behind the best ones, then read the numbers together — reach, followers, free signups — and decide to scale, adjust, or stop. No open-ended commitment.”')

heading(d2,'7 · Payment — give him the choice', size=14)
bullet(d2,[('Option A — milestone release: ',True),('settles in ~7 days, stays on-platform.',False)])
bullet(d2,[('Option B — direct to your US account: ',True),('clears in ~2 days, faster start.',False)])
quote(d2,'“Either works for me. It’s just how fast you want the next cycle moving — if you want it live this week, the direct route is quicker.”')

heading(d2,'Numbers cheat-sheet (memorise these)', size=14)
make_table(d2, [
    ['Item','Number'],
    ['Milestone 1 (release now)','$750'],
    ['Tool stack / month','$235'],
    ['Paid boost (7-day, Standard)','$595'],
    ['One-cycle budget (incl. buffer)','~$915'],
    ['Bonus rate','$5 / free signup'],
    ['Bonus target / payout','100 → $500'],
    ['Scale option','250 → $1,000'],
    ['Pay: milestone','~7 days'],
    ['Pay: direct US account','~2 days'],
], widths=[3.8,2.4], num_cols=[1])

heading(d2,'Likely objections → your answer', size=14)
obj = [
 ('“Why pay $750 if it didn’t go viral?”','It’s for delivered production, not a virality guarantee. Reach is what the new plan fixes — fully budgeted, your call to fund.'),
 ('“Isn’t copying viral posts risky/unoriginal?”','We model structure, not content. Standard practice for every serious creator. Brand stays ours.'),
 ('“Why pay for free signups, not sales?”','That’s the part I control and can guarantee effort against. Sales depend on your funnel. You also keep the audience.'),
 ('“$915 is a lot.”','Most of it is your ad spend going to your reach, plus tools you keep. We can run the Lean tier (~$315 boost) to start smaller.'),
 ('“Can we skip paid boost?”','That’s the exact thing missing in month one. Without it we’d repeat what didn’t work.'),
]
for q,a in obj:
    p=d2.add_paragraph(style='List Bullet'); p.paragraph_format.space_after=Pt(3)
    run(p,q+'  ',size=10.5,bold=True,color=NAVY); run(p,'→ '+a,size=10.5,color=DARK)

heading(d2,'After the call — confirm in writing', size=14)
body(d2,'Send a 4-line recap: milestone released (Y/N + method) · budget approved (tier) · bonus terms ($5/signup, target) · start date.', size=10.5)

d2.save('/home/user/Jejeola11/strategy-reset-v2/Strategy_Reset_v2_Call_Run_Sheet.docx')
print('RUN-SHEET saved')
